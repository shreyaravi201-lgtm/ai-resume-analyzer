import io
from pypdf import PdfReader
from docx import Document

def extract_pdf_text(file_obj) -> str:
    """Extracts text from a PDF file (path or file-like object) page by page."""
    try:
        reader = PdfReader(file_obj)
        extracted_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                extracted_text.append(text)
        
        full_text = "\n".join(extracted_text).strip()
        if not full_text:
            raise ValueError("No extractable text found in this PDF file.")
        return full_text
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file: {str(e)}")

def extract_docx_text(file_obj) -> str:
    """Extracts text from a DOCX file (path or file-like object) paragraphs and tables."""
    try:
        # Document can open a path or a file-like object
        doc = Document(file_obj)
        extracted_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    extracted_text.append(" | ".join(row_text))
                    
        full_text = "\n".join(extracted_text).strip()
        if not full_text:
            raise ValueError("No extractable text found in this DOCX file.")
        return full_text
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")

def parse_resume(file_obj, filename: str) -> str:
    """Dispatches parser based on file extension."""
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_pdf_text(file_obj)
    elif ext == "docx":
        return extract_docx_text(file_obj)
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Please upload a PDF or DOCX file.")
