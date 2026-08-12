import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def create_docx_data_analyst(output_path: str):
    """Creates a synthetic Data Analyst DOCX resume."""
    doc = Document()
    
    doc.add_heading('Alex Johnson', level=0)
    
    doc.add_paragraph('Email: alex.johnson@email.com | Phone: (555) 019-2834 | San Francisco, CA')
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        'Result-driven Data Analyst with 3+ years of experience transforming raw business data '
        'into actionable insights. Skilled in writing SQL queries, performing exploratory data '
        'manipulation in Pandas, and building interactive dashboards to present key business metrics.'
    )
    
    doc.add_heading('Technical Skills', level=1)
    p = doc.add_paragraph()
    p.add_run('Programming: ').bold = True
    p.add_run('Python, SQL, R\n')
    p.add_run('Data & Analytics: ').bold = True
    p.add_run('Pandas, NumPy, Excel, Power BI, Tableau\n')
    p.add_run('Tools & DevOps: ').bold = True
    p.add_run('Git, GitHub, Jupyter Notebook')
    
    doc.add_heading('Work Experience', level=1)
    doc.add_heading('Data Analyst | TechInsights Corp | 2024 - Present', level=2)
    doc.add_paragraph(
        '- Wrote complex SQL joins and aggregations to extract monthly performance reports.\n'
        '- Developed data cleaning pipelines using Pandas and NumPy, reducing preprocessing time by 40%.\n'
        '- Built dynamic Power BI dashboards for executives to monitor sales funnel conversion rates.'
    )
    
    doc.add_heading('Projects', level=1)
    doc.add_heading('Customer Churn Dashboard | Power BI, SQL, Python', level=2)
    doc.add_paragraph(
        'Designed an interactive end-to-end dashboard mapping churn rates. Used Python to cluster '
        'customers, SQL to store the data, and Power BI to render monthly visual charts.'
    )
    
    doc.add_heading('Education & Certifications', level=1)
    doc.add_paragraph('B.S. in Statistics | University of California, Berkeley | 2020 - 2024\n'
                      'Certified Data Analyst Associate (Microsoft PL-300)')
    
    doc.save(output_path)
    print(f"Created DOCX resume: {output_path}")

def create_pdf_ml_engineer(output_path: str):
    """Creates a synthetic Machine Learning Engineer PDF resume using ReportLab."""
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#2B6CB0"),
        spaceAfter=5
    )
    
    h1_style = ParagraphStyle(
        'H1Style',
        parent=styles['Heading2'],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#2D3748"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['BodyText'],
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor("#4A5568")
    )
    
    story = [
        Paragraph("Taylor Vance", title_style),
        Paragraph("taylor.vance@ml-corp.net | +1 (555) 432-1098 | Seattle, WA", body_style),
        Spacer(1, 10),
        
        Paragraph("Professional Summary", h1_style),
        Paragraph(
            "Innovative Machine Learning Engineer specializing in model development, optimization, "
            "and API deployments. Skilled at taking prototype algorithms and scaling them to "
            "production using containerization and clean REST API boundaries.",
            body_style
        ),
        
        Paragraph("Technical Skills", h1_style),
        Paragraph(
            "<b>Languages & Frameworks:</b> Python, Scikit-learn, TensorFlow, PyTorch, FastAPI, SQL<br/>"
            "<b>Data Processing:</b> Pandas, NumPy, Scipy<br/>"
            "<b>MLOps & Cloud:</b> Docker, MLflow, AWS (EC2, S3), Git, GitHub, CI/CD",
            body_style
        ),
        
        Paragraph("Professional Experience", h1_style),
        Paragraph("<b>Machine Learning Engineer | Apex AI | 2024 - Present</b><br/>"
                  "- Built and optimized predictive models using Scikit-learn and XGBoost.<br/>"
                  "- Wrapped models in asynchronous FastAPI services to serve live inferences with under 50ms latency.<br/>"
                  "- Containerized applications using Docker to facilitate deployment to AWS cloud environments.",
                  body_style),
        
        Paragraph("Projects", h1_style),
        Paragraph("<b>End-to-End Model Tracking Pipeline</b><br/>"
                  "Implemented MLflow to track model hyperparameters, validation metrics, and weight checkpoints. "
                  "Automated deployment builds using Git version control and GitHub Actions CI/CD.",
                  body_style),
        
        Paragraph("Education", h1_style),
        Paragraph("M.S. in Computer Science | University of Washington | 2022 - 2024", body_style)
    ]
    
    doc.build(story)
    print(f"Created PDF resume: {output_path}")

def create_pdf_nlp_engineer(output_path: str):
    """Creates a synthetic NLP Engineer PDF resume using ReportLab."""
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#2B6CB0"),
        spaceAfter=5
    )
    
    h1_style = ParagraphStyle(
        'H1Style',
        parent=styles['Heading2'],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#2D3748"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['BodyText'],
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor("#4A5568")
    )
    
    story = [
        Paragraph("Jordan Patel", title_style),
        Paragraph("jordan.patel@language-tech.org | +1 (555) 789-3412 | Boston, MA", body_style),
        Spacer(1, 10),
        
        Paragraph("Professional Summary", h1_style),
        Paragraph(
            "Natural Language Processing (NLP) Engineer with experience designing and applying Transformer-based "
            "architectures. Experienced in textual cleaning, preprocessing pipelines, semantic search, and "
            "working with state-of-the-art Hugging Face models.",
            body_style
        ),
        
        Paragraph("Technical Skills", h1_style),
        Paragraph(
            "<b>NLP & ML:</b> NLP, Transformers, Hugging Face, Deep Learning, PyTorch, Scikit-learn<br/>"
            "<b>Programming & DB:</b> Python, SQL, PostgreSQL, Git, GitHub<br/>"
            "<b>Deployment:</b> Docker, FastAPI",
            body_style
        ),
        
        Paragraph("Professional Experience", h1_style),
        Paragraph("<b>NLP Research Engineer | Language Labs | 2024 - Present</b><br/>"
                  "- Built text processing tokenizers and fine-tuned BERT and GPT models using Hugging Face Transformers.<br/>"
                  "- Implemented semantic search engines on top of PostgreSQL data stores using PyTorch embeddings.<br/>"
                  "- Deployed multi-threaded text classification pipelines in Dockerized FastAPI microservices.",
                  body_style),
        
        Paragraph("Projects", h1_style),
        Paragraph("<b>Conversational QA System</b><br/>"
                  "Engineered QA systems with Hugging Face architectures. Applied extensive textual clean preprocessing "
                  "to eliminate duplicate spacing, HTML tags, and noisy punctuation prior to tokenization.",
                  body_style),
        
        Paragraph("Education", h1_style),
        Paragraph("B.S. in Computer Science & Linguistics | Boston University | 2020 - 2024", body_style)
    ]
    
    doc.build(story)
    print(f"Created PDF resume: {output_path}")

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    create_docx_data_analyst(os.path.join(current_dir, "data_analyst_resume.docx"))
    create_pdf_ml_engineer(os.path.join(current_dir, "ml_engineer_resume.pdf"))
    create_pdf_nlp_engineer(os.path.join(current_dir, "nlp_engineer_resume.pdf"))
